"""
Module 1 — Resume Parser
Extracts clean text from PDF and DOCX files.
Supports: local file path, bytes, or UploadFile (FastAPI).
"""

import io
import re
from pathlib import Path


def parse_resume(file_bytes: bytes, filename: str) -> dict:
    """
    Main entry point.
    Returns: { "text": str, "file_type": str, "page_count": int, "error": str|None }
    """
    ext = Path(filename).suffix.lower()

    if ext == ".pdf":
        return _parse_pdf(file_bytes)
    elif ext in (".docx", ".doc"):
        return _parse_docx(file_bytes)
    else:
        return {"text": "", "file_type": ext, "page_count": 0, "error": f"Unsupported file type: {ext}"}


def _parse_pdf(file_bytes: bytes) -> dict:
    try:
        import fitz  # pymupdf
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        pages = []
        for page in doc:
            pages.append(page.get_text("text"))
        raw = "\n".join(pages)
        return {
            "text": _clean_text(raw),
            "file_type": "pdf",
            "page_count": len(doc),
            "error": None,
        }
    except Exception as e:
        return {"text": "", "file_type": "pdf", "page_count": 0, "error": str(e)}


def _parse_docx(file_bytes: bytes) -> dict:
    try:
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Also pull text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())
        raw = "\n".join(paragraphs)
        return {
            "text": _clean_text(raw),
            "file_type": "docx",
            "page_count": 1,  # DOCX doesn't expose page count easily
            "error": None,
        }
    except Exception as e:
        return {"text": "", "file_type": "docx", "page_count": 0, "error": str(e)}


def _clean_text(text: str) -> str:
    """Remove junk characters, normalize whitespace."""
    # Remove non-printable characters
    text = re.sub(r"[^\x20-\x7E\n]", " ", text)
    # Collapse multiple blank lines into one
    text = re.sub(r"\n{3,}", "\n\n", text)
    # Collapse multiple spaces
    text = re.sub(r"[ \t]{2,}", " ", text)
    return text.strip()
